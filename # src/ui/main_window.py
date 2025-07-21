# src/ui/main_window.py
"""
Tek karakter görünümü - Navigate edilebilir UI - TAM VERSİYON
"""
import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
import os
import sys
from typing import List, Optional
import tempfile
import webbrowser
from pathlib import Path

# Word to HTML conversion için
try:
    from docx import Document
    import mammoth  # Word'den HTML'e dönüştürme için
except ImportError:
    print("mammoth kütüphanesi bulunamadı. pip install python-mammoth ile yükleyin")

# Servis importları
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from services.word_reader import WordReaderService
from services.data_processor import DataProcessorService, TeknikResimKarakteri

class SingleKarakterView(ctk.CTkFrame):
    """Tek karakter görünümü - büyük ve detaylı"""
    
    def __init__(self, parent, on_update_callback=None):
        super().__init__(parent)
        
        self.current_karakter: Optional[TeknikResimKarakteri] = None
        self.on_update_callback = on_update_callback
        self.actual_entry = None
        
        self.setup_ui()
    
    def setup_ui(self):
        """UI'ı oluşturur"""
        # Ana grid ayarları
        self.grid_columnconfigure(1, weight=1)
        
        # Başlık alanı
        self.title_frame = ctk.CTkFrame(self)
        self.title_frame.grid(row=0, column=0, columnspan=2, sticky="ew", padx=20, pady=(20, 10))
        
        self.item_label = ctk.CTkLabel(
            self.title_frame,
            text="Karakter seçilmedi",
            font=ctk.CTkFont(size=24, weight="bold"),
            text_color="white"
        )
        self.item_label.pack(pady=15)
        
        # Ana bilgiler
        info_frame = ctk.CTkFrame(self)
        info_frame.grid(row=1, column=0, columnspan=2, sticky="ew", padx=20, pady=10)
        info_frame.grid_columnconfigure(1, weight=1)
        
        # Dimension
        dim_label = ctk.CTkLabel(info_frame, text="Ölçü:", font=ctk.CTkFont(size=16, weight="bold"))
        dim_label.grid(row=0, column=0, sticky="w", padx=20, pady=10)
        
        self.dim_value = ctk.CTkLabel(
            info_frame, 
            text="-",
            font=ctk.CTkFont(size=16),
            wraplength=300
        )
        self.dim_value.grid(row=0, column=1, sticky="w", padx=20, pady=10)
        
        # Tooling
        tool_label = ctk.CTkLabel(info_frame, text="Ölçüm Aleti:", font=ctk.CTkFont(size=16, weight="bold"))
        tool_label.grid(row=1, column=0, sticky="w", padx=20, pady=10)
        
        self.tool_value = ctk.CTkLabel(
            info_frame, 
            text="-",
            font=ctk.CTkFont(size=16)
        )
        self.tool_value.grid(row=1, column=1, sticky="w", padx=20, pady=10)
        
        # BP Zone
        zone_label = ctk.CTkLabel(info_frame, text="Bölge:", font=ctk.CTkFont(size=16, weight="bold"))
        zone_label.grid(row=2, column=0, sticky="w", padx=20, pady=10)
        
        self.zone_value = ctk.CTkLabel(
            info_frame, 
            text="-",
            font=ctk.CTkFont(size=16)
        )
        self.zone_value.grid(row=2, column=1, sticky="w", padx=20, pady=10)
        
        # Inspection Level
        level_label = ctk.CTkLabel(info_frame, text="Kontrol Seviyesi:", font=ctk.CTkFont(size=16, weight="bold"))
        level_label.grid(row=3, column=0, sticky="w", padx=20, pady=10)
        
        self.level_value = ctk.CTkLabel(
            info_frame, 
            text="-",
            font=ctk.CTkFont(size=16)
        )
        self.level_value.grid(row=3, column=1, sticky="w", padx=20, pady=10)
        
        # Remarks
        remarks_label = ctk.CTkLabel(info_frame, text="Açıklamalar:", font=ctk.CTkFont(size=16, weight="bold"))
        remarks_label.grid(row=4, column=0, sticky="nw", padx=20, pady=10)
        
        self.remarks_value = ctk.CTkLabel(
            info_frame, 
            text="-",
            font=ctk.CTkFont(size=14),
            wraplength=400,
            justify="left"
        )
        self.remarks_value.grid(row=4, column=1, sticky="w", padx=20, pady=10)
        
        # Ölçüm girişi - En önemli kısım
        measurement_frame = ctk.CTkFrame(self)
        measurement_frame.grid(row=2, column=0, columnspan=2, sticky="ew", padx=20, pady=20)
        measurement_frame.grid_columnconfigure(1, weight=1)
        
        # Ölçüm başlığı
        measurement_title = ctk.CTkLabel(
            measurement_frame,
            text="ÖLÇÜM SONUCU",
            font=ctk.CTkFont(size=18, weight="bold"),
            text_color="yellow"
        )
        measurement_title.grid(row=0, column=0, columnspan=3, pady=(15, 10))
        
        # Mevcut değer göstergesi
        current_label = ctk.CTkLabel(
            measurement_frame,
            text="Mevcut Değer:",
            font=ctk.CTkFont(size=14, weight="bold")
        )
        current_label.grid(row=1, column=0, sticky="w", padx=20, pady=5)
        
        self.current_value_label = ctk.CTkLabel(
            measurement_frame,
            text="Henüz ölçüm yapılmadı",
            font=ctk.CTkFont(size=14),
            text_color="orange"
        )
        self.current_value_label.grid(row=1, column=1, sticky="w", padx=10, pady=5)
        
        # Yeni değer girişi
        new_label = ctk.CTkLabel(
            measurement_frame,
            text="Yeni Ölçüm:",
            font=ctk.CTkFont(size=14, weight="bold")
        )
        new_label.grid(row=2, column=0, sticky="w", padx=20, pady=10)
        
        self.actual_entry = ctk.CTkEntry(
            measurement_frame,
            placeholder_text="Ölçüm değerini girin (örn: 25.48)",
            width=200,
            height=35,
            font=ctk.CTkFont(size=14)
        )
        self.actual_entry.grid(row=2, column=1, sticky="w", padx=10, pady=10)
        
        # Kaydet butonu
        self.save_button = ctk.CTkButton(
            measurement_frame,
            text="Kaydet",
            command=self.save_measurement,
            width=100,
            height=35,
            font=ctk.CTkFont(size=14, weight="bold")
        )
        self.save_button.grid(row=2, column=2, sticky="w", padx=10, pady=10)
        
        # Temizle butonu
        clear_button = ctk.CTkButton(
            measurement_frame,
            text="Temizle",
            command=self.clear_measurement,
            width=80,
            height=35,
            font=ctk.CTkFont(size=12),
            fg_color="gray"
        )
        clear_button.grid(row=3, column=1, sticky="w", padx=10, pady=5)
        
        # Status mesajı
        self.status_label = ctk.CTkLabel(
            measurement_frame,
            text="",
            font=ctk.CTkFont(size=12)
        )
        self.status_label.grid(row=4, column=0, columnspan=3, pady=10)
        
        # Enter tuşu ile kaydetme
        self.actual_entry.bind('<Return>', lambda e: self.save_measurement())
    
    def load_karakter(self, karakter: TeknikResimKarakteri):
        """Karakteri yükler ve gösterir"""
        self.current_karakter = karakter
        
        # Bilgileri güncelle
        self.item_label.configure(text=f"Item: {karakter.item_no}")
        self.dim_value.configure(text=karakter.dimension)
        self.tool_value.configure(text=karakter.tooling)
        self.zone_value.configure(text=karakter.bp_zone or "Belirtilmemiş")
        self.level_value.configure(text=karakter.inspection_level or "100%")
        self.remarks_value.configure(text=karakter.remarks or "Açıklama yok")
        
        # Mevcut ölçüm değerini göster
        if karakter.actual:
            self.current_value_label.configure(
                text=f"{karakter.actual}",
                text_color="green"
            )
            # Entry'e de yerleştir
            self.actual_entry.delete(0, tk.END)
            self.actual_entry.insert(0, str(karakter.actual))
        else:
            self.current_value_label.configure(
                text="Henüz ölçüm yapılmadı",
                text_color="orange"
            )
            self.actual_entry.delete(0, tk.END)
        
        # Status'u temizle
        self.status_label.configure(text="")
        
        # Entry'e focus ver
        self.actual_entry.focus()
    
    def save_measurement(self):
        """Ölçümü kaydeder"""
        if not self.current_karakter:
            return
        
        try:
            new_value = self.actual_entry.get().strip()
            
            if new_value == "":
                self.status_label.configure(text="⚠ Değer boş bırakılamaz", text_color="orange")
                return
            
            # Virgülü noktaya çevir
            new_value = new_value.replace(',', '.')
            
            # Sayı kontrolü (isteğe bağlı)
            try:
                float(new_value)  # Sayı mı kontrol et
                self.current_karakter.actual = new_value
                self.status_label.configure(text="✓ Ölçüm kaydedildi!", text_color="green")
            except ValueError:
                # Sayı değilse de kabul et
                self.current_karakter.actual = new_value
                self.status_label.configure(text="✓ Kaydedildi (metin değer)", text_color="green")
            
            # Mevcut değer göstergesini güncelle
            self.current_value_label.configure(
                text=f"{new_value}",
                text_color="green"
            )
            
            # Callback çağır
            if self.on_update_callback:
                self.on_update_callback(self.current_karakter)
            
        except Exception as e:
            self.status_label.configure(text=f"Hata: {str(e)}", text_color="red")
    
    def clear_measurement(self):
        """Ölçümü temizler"""
        if self.current_karakter:
            self.current_karakter.actual = None
            self.actual_entry.delete(0, tk.END)
            self.current_value_label.configure(
                text="Henüz ölçüm yapılmadı",
                text_color="orange"
            )
            self.status_label.configure(text="Ölçüm temizlendi", text_color="gray")
            
            if self.on_update_callback:
                self.on_update_callback(self.current_karakter)


class NavigationPanel(ctk.CTkFrame):
    """Navigasyon paneli - önceki/sonraki butonları"""
    
    def __init__(self, parent, on_navigate_callback=None):
        super().__init__(parent)
        
        self.on_navigate_callback = on_navigate_callback
        self.current_index = 0
        self.total_count = 0
        
        self.setup_ui()
    
    def setup_ui(self):
        """Navigasyon UI'ı"""
        # Önceki butonu
        self.prev_button = ctk.CTkButton(
            self,
            text="◀ Önceki",
            command=self.go_previous,
            width=120,
            height=40,
            font=ctk.CTkFont(size=14, weight="bold"),
            state="disabled"
        )
        self.prev_button.pack(side="left", padx=20, pady=15)
        
        # Pozisyon göstergesi
        self.position_label = ctk.CTkLabel(
            self,
            text="0 / 0",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        self.position_label.pack(side="left", padx=30, pady=15)
        
        # Sonraki butonu
        self.next_button = ctk.CTkButton(
            self,
            text="Sonraki ▶",
            command=self.go_next,
            width=120,
            height=40,
            font=ctk.CTkFont(size=14, weight="bold"),
            state="disabled"
        )
        self.next_button.pack(side="left", padx=20, pady=15)
        
        # Progress bar
        self.progress = ctk.CTkProgressBar(self, width=200)
        self.progress.pack(side="right", padx=20, pady=15)
        self.progress.set(0)
    
    def update_navigation(self, current_index: int, total_count: int):
        """Navigasyon durumunu günceller"""
        self.current_index = current_index
        self.total_count = total_count
        
        # Pozisyon etiketi
        self.position_label.configure(text=f"{current_index + 1} / {total_count}")
        
        # Buton durumları
        self.prev_button.configure(state="normal" if current_index > 0 else "disabled")
        self.next_button.configure(state="normal" if current_index < total_count - 1 else "disabled")
        
        # Progress bar
        if total_count > 0:
            progress = (current_index + 1) / total_count
            self.progress.set(progress)
    
    def go_previous(self):
        """Önceki karaktere git"""
        if self.current_index > 0 and self.on_navigate_callback:
            self.on_navigate_callback(self.current_index - 1)
    
    def go_next(self):
        """Sonraki karaktere git"""
        if self.current_index < self.total_count - 1 and self.on_navigate_callback:
            self.on_navigate_callback(self.current_index + 1)


class DocumentViewer(ctk.CTkFrame):
    """Word dokümanını görüntülemek için panel"""
    
    def __init__(self, parent):
        super().__init__(parent)
        
        self.current_html_file = None
        self.setup_ui()
    
    def setup_ui(self):
        """Doküman görüntüleyici UI"""
        # Başlık
        title_label = ctk.CTkLabel(
            self,
            text="Word Dokümanı",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        title_label.pack(pady=(10, 5))
        
        # Doküman yükleme butonu
        self.load_button = ctk.CTkButton(
            self,
            text="Dokümanı HTML'de Aç",
            command=self.open_in_browser,
            state="disabled"
        )
        self.load_button.pack(pady=5)
        
        # Text area - HTML preview için
        self.text_area = ctk.CTkTextbox(
            self,
            wrap="word",
            font=ctk.CTkFont(size=11)
        )
        self.text_area.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Başlangıç mesajı
        self.text_area.insert("1.0", "Word dosyası yüklendiğinde doküman içeriği burada görünecek.\n\nTam görünüm için 'Dokümanı HTML'de Aç' butonunu kullanın.")
        self.text_area.configure(state="disabled")
    
    def load_document(self, file_path: str):
        """Word dokümanını yükler ve gösterir"""
        try:
            print(f"Doküman yükleniyor: {file_path}")
            
            # Text area'yı temizle
            self.text_area.configure(state="normal")
            self.text_area.delete("1.0", tk.END)
            self.text_area.insert("1.0", "Doküman yükleniyor...\n")
            self.text_area.configure(state="disabled")
            
            # Word dokümanını oku
            with open(file_path, "rb") as docx_file:
                # mammoth ile HTML'e çevir
                try:
                    result = mammoth.convert_to_html(docx_file)
                    html_content = result.html
                    
                    # HTML dosyası oluştur
                    self.create_html_file(html_content, file_path)
                    
                    # Text area'da basit preview göster
                    self.show_text_preview(file_path)
                    
                    # HTML açma butonunu aktif et
                    self.load_button.configure(state="normal")
                    
                    print("✓ Doküman başarıyla yüklendi")
                    
                except ImportError:
                    # mammoth yoksa basit text extraction
                    self.show_simple_text(file_path)
                    
        except Exception as e:
            error_msg = f"Doküman yükleme hatası: {str(e)}"
            print(error_msg)
            
            self.text_area.configure(state="normal")
            self.text_area.delete("1.0", tk.END)
            self.text_area.insert("1.0", f"HATA: {error_msg}")
            self.text_area.configure(state="disabled")
    
    def create_html_file(self, html_content: str, original_file: str):
        """HTML dosyası oluşturur"""
        try:
            # Geçici HTML dosyası oluştur
            temp_dir = tempfile.gettempdir()
            file_name = Path(original_file).stem
            html_file = os.path.join(temp_dir, f"{file_name}_preview.html")
            
            # CSS ile biçimlendirme
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <title>Word Dokümanı Preview</title>
                <style>
                    body {{ 
                        font-family: Arial, sans-serif; 
                        margin: 20px; 
                        line-height: 1.6;
                    }}
                    table {{ 
                        border-collapse: collapse; 
                        width: 100%; 
                        margin: 10px 0;
                    }}
                    th, td {{ 
                        border: 1px solid #ddd; 
                        padding: 8px; 
                        text-align: left;
                    }}
                    th {{ 
                        background-color: #f2f2f2; 
                        font-weight: bold;
                    }}
                </style>
            </head>
            <body>
                <h1>Word Dokümanı - {file_name}</h1>
                {html_content}
            </body>
            </html>
            """
            
            with open(html_file, 'w', encoding='utf-8') as f:
                f.write(full_html)
            
            self.current_html_file = html_file
            print(f"✓ HTML dosyası oluşturuldu: {html_file}")
            
        except Exception as e:
            print(f"HTML dosyası oluşturma hatası: {e}")
    
    def show_text_preview(self, file_path: str):
        """Text area'da basit preview göster"""
        try:
            # Word dokümanından text çıkar
            doc = Document(file_path)
            
            text_content = "=== WORD DOKÜMAN İÇERİĞİ ===\n\n"
            
            # Paragrafları ekle
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content += para.text + "\n"
            
            # Tabloları ekle
            text_content += "\n=== TABLOLAR ===\n"
            for i, table in enumerate(doc.tables):
                text_content += f"\n--- Tablo {i+1} ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    text_content += row_text + "\n"
            
            # Text area'yı güncelle
            self.text_area.configure(state="normal")
            self.text_area.delete("1.0", tk.END)
            self.text_area.insert("1.0", text_content)
            self.text_area.configure(state="disabled")
            
        except Exception as e:
            error_msg = f"Text preview hatası: {str(e)}"
            self.text_area.configure(state="normal")
            self.text_area.delete("1.0", tk.END)
            self.text_area.insert("1.0", error_msg)
            self.text_area.configure(state="disabled")
    
    def show_simple_text(self, file_path: str):
        """mammoth yoksa basit text gösterimi"""
        try:
            doc = Document(file_path)
            
            content = "mammoth kütüphanesi bulunamadı. Basit text görünümü:\n\n"
            
            for para in doc.paragraphs:
                if para.text.strip():
                    content += para.text + "\n\n"
            
            self.text_area.configure(state="normal")
            self.text_area.delete("1.0", tk.END)
            self.text_area.insert("1.0", content)
            self.text_area.configure(state="disabled")
            
        except Exception as e:
            print(f"Basit text gösterimi hatası: {e}")
    
    def open_in_browser(self):
        """HTML dosyasını tarayıcıda aç"""
        if self.current_html_file and os.path.exists(self.current_html_file):
            try:
                webbrowser.open(f'file://{self.current_html_file}')
                print(f"HTML dosyası tarayıcıda açıldı: {self.current_html_file}")
            except Exception as e:
                messagebox.showerror("Hata", f"Tarayıcıda açılamadı: {str(e)}")
        else:
            messagebox.showwarning("Uyarı", "Önce bir doküman yükleyin!")


class NavigableMainWindow(ctk.CTk):
    """Navigate edilebilir ana pencere"""
    
    def __init__(self):
        super().__init__()
        
        # Pencere ayarları
        self.title("Teknik Resim Karakter Okuyucu - Navigate Edilebilir")
        self.geometry("1400x900")
        
        # Servisler
        self.word_service = WordReaderService()
        self.data_service = DataProcessorService()
        
        # Veri
        self.karakterler: List[TeknikResimKarakteri] = []
        self.current_index = 0
        self.current_file_path: Optional[str] = None
        
        self.setup_ui()
        
        # Klavye shortcut'ları
        self.bind('<Left>', lambda e: self.navigate_to(self.current_index - 1))
        self.bind('<Right>', lambda e: self.navigate_to(self.current_index + 1))
        self.focus_set()  # Klavye focus'u için
    
    def setup_ui(self):
        """Ana UI'ı oluşturur"""
        # Grid ayarları - 2 eşit kolon
        self.grid_columnconfigure(0, weight=1)  # Sol panel
        self.grid_columnconfigure(1, weight=1)  # Sağ panel
        self.grid_rowconfigure(1, weight=1)     # Ana içerik
        
        # Üst panel - Dosya işlemleri
        self.create_top_panel()
        
        # Sol panel - Tek karakter görünümü
        self.create_left_panel()
        
        # Sağ panel - Doküman görüntüleyici
        self.create_right_panel()
        
        # Alt panel - İstatistikler
        self.create_bottom_panel()
    
    def create_top_panel(self):
        """Üst panel - dosya seçme"""
        top_frame = ctk.CTkFrame(self)
        top_frame.grid(row=0, column=0, columnspan=2, sticky="ew", padx=10, pady=10)
        
        # Dosya seç butonu
        self.file_button = ctk.CTkButton(
            top_frame,
            text="Word Dosyası Seç",
            command=self.select_file,
            height=40,
            font=ctk.CTkFont(size=14)
        )
        self.file_button.pack(side="left", padx=10, pady=10)
        
        # Dosya yolu gösterici
        self.file_path_label = ctk.CTkLabel(
            top_frame,
            text="Dosya seçilmedi",
            font=ctk.CTkFont(size=12)
        )
        self.file_path_label.pack(side="left", padx=10, pady=10)
        
        # İşle butonu
        self.process_button = ctk.CTkButton(
            top_frame,
            text="Dosyayı Yükle",
            command=self.process_file,
            height=40,
            font=ctk.CTkFont(size=14),
            state="disabled"
        )
        self.process_button.pack(side="right", padx=10, pady=10)
    
    def create_left_panel(self):
        """Sol panel - tek karakter görünümü"""
        left_frame = ctk.CTkFrame(self)
        left_frame.grid(row=1, column=0, sticky="nsew", padx=(10, 5), pady=10)
        left_frame.grid_rowconfigure(0, weight=1)
        left_frame.grid_columnconfigure(0, weight=1)
        
        # Karakter görünümü
        self.karakter_view = SingleKarakterView(
            left_frame,
            on_update_callback=self.on_karakter_updated
        )
        self.karakter_view.grid(row=0, column=0, sticky="nsew", padx=10, pady=(10, 5))
        
        # Navigasyon paneli
        self.navigation_panel = NavigationPanel(
            left_frame,
            on_navigate_callback=self.navigate_to
        )
        self.navigation_panel.grid(row=1, column=0, sticky="ew", padx=10, pady=(5, 10))
    
    def create_right_panel(self):
        """Sağ panel - doküman görüntüleyici"""
        right_frame = ctk.CTkFrame(self)
        right_frame.grid(row=1, column=1, sticky="nsew", padx=(5, 10), pady=10)
        right_frame.grid_rowconfigure(0, weight=1)
        right_frame.grid_columnconfigure(0, weight=1)
        
        # Doküman görüntüleyici
        self.document_viewer = DocumentViewer(right_frame)
        self.document_viewer.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
    
    def create_bottom_panel(self):
        """Alt panel - istatistikler ve kaydetme"""
        bottom_frame = ctk.CTkFrame(self)
        bottom_frame.grid(row=2, column=0, columnspan=2, sticky="ew", padx=10, pady=10)
        
        # İstatistik labelı
        self.stats_label = ctk.CTkLabel(
            bottom_frame,
            text="İstatistikler burada görünecek",
            font=ctk.CTkFont(size=12)
        )
        self.stats_label.pack(side="left", padx=10, pady=10)
        
        # Kaydetme butonları
        save_frame = ctk.CTkFrame(bottom_frame, fg_color="transparent")
        save_frame.pack(side="right", padx=10, pady=5)
        
        export_button = ctk.CTkButton(
            save_frame,
            text="Excel'e Aktar",
            command=self.export_to_excel,
            height=30
        )
        export_button.pack(side="right", padx=5)
    
    def select_file(self):
        """Word dosyası seçme"""
        file_path = filedialog.askopenfilename(
            title="Word Dosyası Seçin",
            filetypes=[
                ("Word Dosyaları", "*.docx *.doc"),
                ("Tüm Dosyalar", "*.*")
            ]
        )
        
        if file_path:
            self.current_file_path = file_path
            file_name = os.path.basename(file_path)
            self.file_path_label.configure(text=f"Seçilen: {file_name}")
            self.process_button.configure(state="normal")
    
    def process_file(self):
        """Dosyayı işler - hem karakterleri hem dokümanı yükler"""
        if not self.current_file_path:
            messagebox.showerror("Hata", "Önce bir dosya seçin!")
            return
        
        try:
            self.file_path_label.configure(text="İşleniyor...")
            self.update()
            
            # 1. Word dosyasını karakterler için işle
            if not self.word_service.load_document(self.current_file_path):
                messagebox.showerror("Hata", "Word dosyası yüklenemedi!")
                return
            
            tables = self.word_service.extract_tables()
            
            if not tables:
                messagebox.showwarning("Uyarı", "Dokümanda tablo bulunamadı!")
                return
            
            self.karakterler = self.data_service.process_dataframe(tables[0])
            
            if not self.karakterler:
                messagebox.showwarning("Uyarı", "Geçerli karakter bulunamadı!")
                return
            
            # 2. Dokümanı sağ panelde göster
            self.document_viewer.load_document(self.current_file_path)
            
            # 3. İlk karakteri göster
            self.current_index = 0
            self.show_current_karakter()
            self.update_navigation()
            self.update_stats()
            
            file_name = os.path.basename(self.current_file_path)
            self.file_path_label.configure(text=f"✓ Yüklendi: {file_name}")
            
            messagebox.showinfo("Başarılı", f"{len(self.karakterler)} karakter yüklendi!\n\nOk tuşları ile navigate edebilirsiniz.")
            
        except Exception as e:
            messagebox.showerror("Hata", f"İşleme hatası:\n{str(e)}")
            print(f"İşleme hatası: {e}")
    
    def show_current_karakter(self):
        """Mevcut karakteri gösterir"""
        if 0 <= self.current_index < len(self.karakterler):
            karakter = self.karakterler[self.current_index]
            self.karakter_view.load_karakter(karakter)
            print(f"Karakter gösteriliyor: {self.current_index + 1}/{len(self.karakterler)} - {karakter.item_no}")
    
    def navigate_to(self, new_index: int):
        """Belirtilen indekse navigate eder"""
        if not self.karakterler:
            return
        
        if 0 <= new_index < len(self.karakterler):
            self.current_index = new_index
            self.show_current_karakter()
            self.update_navigation()
            print(f"Navigate: {self.current_index + 1}/{len(self.karakterler)}")
    
    def update_navigation(self):
        """Navigasyon durumunu günceller"""
        if self.karakterler:
            self.navigation_panel.update_navigation(self.current_index, len(self.karakterler))
    
    def on_karakter_updated(self, karakter: TeknikResimKarakteri):
        """Karakter güncellendiğinde çağrılır"""
        print(f"Karakter güncellendi: {karakter.item_no} = {karakter.actual}")
        self.update_stats()
    
    def update_stats(self):
        """İstatistikleri günceller"""
        if not self.karakterler:
            return
        
        # Ölçüm durumu
        total = len(self.karakterler)
        measured = len([k for k in self.karakterler if k.actual])
        unmeasured = total - measured
        
        # Mevcut karakter bilgisi
        current_info = f"Şu an: {self.current_index + 1}/{total}"
        
        # İstatistik metni
        stats_text = f"{current_info} | Toplam: {total} | Ölçülen: {measured} | Bekleyen: {unmeasured}"
        
        if total > 0:
            percentage = (measured / total) * 100
            stats_text += f" | Tamamlanan: %{percentage:.1f}"
        
        self.stats_label.configure(text=stats_text)
    
    def export_to_excel(self):
        """Sonuçları Excel'e aktarır"""
        if not self.karakterler:
            messagebox.showwarning("Uyarı", "Önce veri yükleyin!")
            return
        
        try:
            # Excel dosyası yolu seç
            file_path = filedialog.asksaveasfilename(
                title="Excel Dosyası Kaydet",
                defaultextension=".xlsx",
                filetypes=[("Excel Dosyaları", "*.xlsx"), ("Tüm Dosyalar", "*.*")]
            )
            
            if file_path:
                import pandas as pd
                
                # DataFrame oluştur
                data = []
                for karakter in self.karakterler:
                    data.append({
                        'Item No': karakter.item_no,
                        'Dimension': karakter.dimension,
                        'Tooling': karakter.tooling,
                        'BP Zone': karakter.bp_zone,
                        'Remarks': karakter.remarks,
                        'Inspection Level': karakter.inspection_level,
                        'Actual': karakter.actual
                    })
                
                df = pd.DataFrame(data)
                df.to_excel(file_path, index=False)
                
                messagebox.showinfo("Başarılı", f"Veriler Excel'e aktarıldı:\n{file_path}")
                
        except Exception as e:
            messagebox.showerror("Hata", f"Excel aktarım hatası:\n{str(e)}")


# Ana çalıştırma
if __name__ == "__main__":
    # CustomTkinter tema ayarları
    ctk.set_appearance_mode("dark")
    ctk.set_default_color_theme("blue")
    
    # Uygulamayı başlat
    app = NavigableMainWindow()
    app.mainloop()
